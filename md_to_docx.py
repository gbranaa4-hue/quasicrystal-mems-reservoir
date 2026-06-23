"""Convert INVESTIGATION_REPORT.md to a formatted .docx (no pandoc/node available)."""
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import sys
SRC = sys.argv[1] if len(sys.argv) > 1 else "INVESTIGATION_REPORT.md"
DST = sys.argv[2] if len(sys.argv) > 2 else SRC.rsplit(".", 1)[0] + ".docx"

with open(SRC, encoding="utf-8") as f:
    lines = f.read().splitlines()

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def add_heading(text, level):
    doc.add_heading(text.strip("# ").strip(), level=level)

def strip_md_links(text):
    return re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)

def add_para(text, bold_whole=False, italic_whole=False):
    text = strip_md_links(text)
    p = doc.add_paragraph()
    # handle inline **bold** segments
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        if italic_whole:
            run.italic = True
        if bold_whole:
            run.bold = True
    return p

def parse_table(block_lines):
    rows = [l.strip() for l in block_lines if l.strip().startswith("|")]
    rows = [r for r in rows if not re.match(r"^\|[\s\-:|]+\|$", r)]
    grid = []
    for r in rows:
        cells = [c.strip() for c in r.strip("|").split("|")]
        grid.append(cells)
    return grid

i = 0
n = len(lines)
in_code_block = False
code_buffer = []

while i < n:
    line = lines[i]

    if line.strip().startswith("```"):
        if not in_code_block:
            in_code_block = True
            code_buffer = []
        else:
            in_code_block = False
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_buffer))
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        i += 1
        continue

    if in_code_block:
        code_buffer.append(line)
        i += 1
        continue

    if line.strip() == "---":
        doc.add_paragraph()
        i += 1
        continue

    if line.startswith("# "):
        add_heading(line, 0)
        i += 1
        continue
    if line.startswith("## "):
        add_heading(line, 1)
        i += 1
        continue
    if line.startswith("### "):
        add_heading(line, 2)
        i += 1
        continue

    if line.strip().startswith("|"):
        block = []
        while i < n and lines[i].strip().startswith("|"):
            block.append(lines[i])
            i += 1
        grid = parse_table(block)
        if grid:
            ncols = len(grid[0])
            table = doc.add_table(rows=0, cols=ncols)
            table.style = "Light Grid Accent 1"
            for r_idx, row in enumerate(grid):
                cells = table.add_row().cells
                for c_idx, val in enumerate(row[:ncols]):
                    cells[c_idx].text = val.replace("**", "")
                    if r_idx == 0:
                        for p in cells[c_idx].paragraphs:
                            for run in p.runs:
                                run.bold = True
            doc.add_paragraph()
        continue

    if re.match(r"^\d+\.\s", line.strip()):
        text = re.sub(r"^\d+\.\s", "", line.strip())
        p = doc.add_paragraph(style="List Number")
        parts = re.split(r"(\*\*[^*]+\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)
        i += 1
        continue

    if line.strip().startswith("- "):
        text = line.strip()[2:]
        p = doc.add_paragraph(style="List Bullet")
        parts = re.split(r"(\*\*[^*]+\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)
        i += 1
        continue

    if line.strip().startswith(">"):
        p = doc.add_paragraph()
        run = p.add_run(line.strip().lstrip("> ").strip())
        run.italic = True
        i += 1
        continue

    if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
        add_para(line.strip().strip("*"), italic_whole=True)
        i += 1
        continue

    if line.strip() == "":
        i += 1
        continue

    add_para(line)
    i += 1

doc.save(DST)
print(f"Saved {DST}")
