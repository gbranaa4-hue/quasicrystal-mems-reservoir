#!/usr/bin/env python3
"""Generate a readable Word document explaining the selection-rule theorem."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- page + base style ----
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1.0))
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

BLUE = RGBColor(0x1F, 0x3B, 0x6E)


def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = BLUE
    p.paragraph_format.space_after = Pt(2)


def subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.italic = True; r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.space_after = Pt(16)


def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)


def body(text, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.italic = italic
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def boxed(text):
    """A visually set-apart statement (centered, bold, blue)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(10)


# ============================================================ TITLE
title("The Symmetry Selection Rule for Even-Order Computation")
subtitle("What it is, what it means, and how I arrived at it")

# ============================================================ 1
h1("1.  The result, in one paragraph")
body(
    "On a resonator with the symmetry of a square, every vibrational mode has a "
    "number attached to it — the integral ∫φ³ over the surface "
    "— that measures whether that mode can perform even-order (product-making) "
    "nonlinear computation. The result is this: on a perfectly square-symmetric "
    "structure, that number is forced to be exactly zero for every mode except the "
    "small handful that are fully symmetric. Seven-eighths of the modes are "
    "“silenced” by symmetry alone, before any calculation. Breaking the "
    "symmetry — with an aperiodic (quasicrystal) hole pattern, or simply by "
    "using a non-square cavity — revives the silenced modes. The vanishing is "
    "not approximate and not statistical: it is exact, provable with pen and paper "
    "from group theory, and confirmed on real finite-element modes to about one "
    "part in a billion.")

# ============================================================ 2
h1("2.  The theorem, built from the ground up")

h2("2.1  Start in one dimension")
body(
    "One fact from calculus contains the whole idea: the integral of an odd "
    "function over a symmetric interval is exactly zero. The area on the left "
    "cancels the area on the right, point for point — symmetry forces it, no "
    "computation required.")
body(
    "A symmetric vibrating object has modes of two kinds: symmetric (mirror-"
    "identical across the center) and antisymmetric (the shape flips sign across "
    "the center). Now take a mode, cube it, and integrate. Symmetry decides the "
    "outcome:")
bullet("If the mode is symmetric, its cube is symmetric → the integral is generally nonzero.")
bullet("If the mode is antisymmetric, its cube is still antisymmetric (odd × odd × odd = odd) → the integral is exactly zero.")
body(
    "That is the entire theorem in miniature. The cube of an antisymmetric mode "
    "cancels itself over the symmetric domain. Half the modes are forced to "
    "∫φ³ = 0 by nothing but their symmetry.")

h2("2.2  Why ∫φ³ is the quantity that matters")
body(
    "∫φ³ is the strength of a mode’s product-making "
    "nonlinearity — the term that lets it multiply two signals together to "
    "compute something like u×u. If ∫φ³ = 0, that mode cannot "
    "contribute to forming products at all. It is not weak at even-order "
    "computation; it is mute. That is the bridge from “an integral vanishes” "
    "to “this mode is computationally silenced.”")

h2("2.3  The square and its eight symmetries")
body(
    "A square has more symmetry than a line: flip left–right, flip up–down, "
    "flip across each diagonal, and rotate by 90°, 180°, 270° — "
    "eight operations in all. Mathematicians call this group D4. Because there are "
    "more symmetries, modes can no longer be just “even” or “odd”; "
    "they fall into five symmetry types (the “irreducible representations”):")
bullet("A1 — unchanged by every operation: the fully symmetric type (like a dome).")
bullet("A2, B1, B2 — flip sign under some operations (different antisymmetric patterns).")
bullet("E — a two-dimensional type whose modes come in degenerate pairs that the rotations mix together.")

h2("2.4  Only the fully symmetric modes survive")
body(
    "The same odd/even logic generalizes exactly: ∫φ³ survives only "
    "if the mode’s type, when cubed, still contains a fully-symmetric (A1) "
    "piece — because integrating over the symmetric shape keeps only the "
    "fully-symmetric part and annihilates everything else. Checking each type:")
bullet("A1 cubed = A1. Contains the symmetric piece → ∫φ³ may be nonzero. (These are the live modes.)")
bullet("A2, B1, B2 cubed = back to A2, B1, B2 — each still carries a sign-flip → ∫φ³ = 0.")
bullet("E cubed decomposes as Sym³(E) = 2E: two copies of the E type and zero copies of A1 → ∫φ³ = 0.")
body(
    "So only the A1 (fully symmetric) modes can carry the nonlinearity, and those "
    "are only about one-eighth of all modes. The other seven-eighths are forced to "
    "exactly zero. That 7/8 is where the measured “88% of modes silenced” "
    "came from — it was never a coincidence; it was 7/8 wearing a disguise.")
boxed("∫φ³ = 0  for every mode except the fully symmetric ones.")

h2("2.5  What “confirmed to one part in a billion” means")
body(
    "The theorem predicts these integrals are exactly zero — a perfect "
    "cancellation. When computed on the actual finite-element mode shapes (real "
    "numbers from a simulation), the forbidden ones came back at about 10⁻⁹, "
    "while the allowed modes came back at about 0.4. That 10⁻⁹ is not "
    "“small”; it is zero to the precision of the computer’s arithmetic. "
    "The separation between forbidden and allowed is a factor of roughly a hundred "
    "million — a hard wall, not a soft trend. That is what makes it a proof and "
    "not a vibe: the vanishing is predicted in advance from symmetry, derivable by "
    "hand, and then confirmed to the limit of the arithmetic.")

# ============================================================ 3
h1("3.  What it means")

h2("3.1  The physical statement")
body(
    "Symmetry decides which modes of a physical resonator are allowed to do "
    "even-order nonlinear computation. A perfectly symmetric structure is blind to "
    "even-order processing in all but its fully-symmetric modes. Breaking the "
    "symmetry switches the rest back on. Aperiodicity does not add richness so much "
    "as it removes a prohibition.")

h2("3.2  The portable principle")
body(
    "Nothing in the argument is specific to plates, holes, or MEMS. It needs only a "
    "physical system whose computation is carried by its modes, an even-order "
    "nonlinearity whose modal strength is ∫φ³, and a point symmetry. "
    "The same rule should therefore apply to photonic cavities, spin-wave devices, "
    "mechanical metamaterials, and electronic resonator arrays. As a design "
    "principle it is compact: to get even-order computation from a modal physical "
    "computer, break the substrate’s point symmetry — and the benefit "
    "grows with how much of the symmetry you break.")

h2("3.3  The honest scope — what is proven and what is not")
body("Proven, exactly:")
bullet("The selection rule itself: ∫φ³ = 0 for all non-symmetric modes, by group theory, confirmed numerically.")
bullet("That breaking the symmetry (by hole pattern OR by cavity shape) revives those modes and yields a measurable even-order advantage in a controlled reservoir model.")
bullet("That the popular intuition — “richer/aperiodic structure computes better” — is false here: spectral and coupling richness are computationally inert. The operative variable is symmetry, not aperiodicity.")
body("Not proven — and stated plainly:")
bullet("That a fabricated device would show a useful advantage (everything here is simulation with a modeled nonlinearity).")
bullet("That the effect is large or practical: it is modest in magnitude and confined to shallow, even-order tasks; deep-memory computation is walled off by a conservation law that no amount of tuning moves.")
body(
    "The result is therefore an exact theorem with a real but bounded computational "
    "consequence. Holding both of those at once — the rule is exact; its payoff "
    "is modest — is the honest description.")

# ============================================================ 4
h1("4.  How I formed it")
body(
    "The result did not fall from the sky, and it was not a master plan. It was a "
    "recombination of pieces I already had, fused by one new move and then carved "
    "down by hard testing.")

h2("4.1  Four parts carried in from earlier work")
bullet(
    "The finite-element plate engine, from my MEMS quasicrystal study. The same "
    "solver — with its homogenized perforation, its de Bruijn quasicrystal "
    "geometry, and its periodic-grid comparison — is the literal foundation. "
    "The whole result is built on code I had written for a different purpose "
    "(frequency filtering).")
bullet(
    "The quasicrystal-versus-periodic question, also from that study. Its central "
    "subject — what does aperiodic perforation do that periodic does not — "
    "became the computing question, just pointed at computation instead of "
    "frequency response.")
bullet(
    "The instinct that a resonator can compute, not only filter — from my "
    "work on resonate-and-fire elements. That is what made “use these plates as "
    "a reservoir” a natural question rather than a random one. The neuromorphic "
    "interest supplied the motivation; the plate model supplied the substrate.")
bullet(
    "The honesty discipline — signal must beat noise, separate what is real "
    "from what is assumed, keep asking ‘does it actually work’ — "
    "carried over from reviewing the paper and stress-testing the filter-bank idea. "
    "It became the controls, ablations, and cross-validation that make the result "
    "trustworthy rather than hopeful.")

h2("4.2  The one new move")
body(
    "Everything above was borrowed. The new step was small but decisive. The "
    "filtering work only ever needed the mode frequencies; it discarded the mode "
    "shapes. The computing question forced me to keep the shapes. Once I had them "
    "and projected a nonlinearity onto them, a quantity appeared that filtering "
    "never had a reason to look at — ∫φ³, the strength of the "
    "product-making nonlinearity. Chasing why that quantity favored the quasicrystal "
    "led to its distribution (most modes near zero), and recognizing that pattern as "
    "a symmetry signature connected it to the group theory of the square — and "
    "that is what turned an observation into a theorem.")

h2("4.3  The honest culling")
body(
    "The earlier projects did not point straight at this. They supplied the "
    "materials and the question; the testing then removed everything that was not "
    "real. The hope that a quasicrystal is simply “richer” and therefore "
    "computes better: tested and discarded. The dream of immediate practical value: "
    "bounded. Later extensions — toward a language model, an all-optical "
    "computer, bandgap-filtered modes — each tested and either refuted or shown "
    "to collapse back into the same one mechanism. The symmetry rule is what was "
    "left standing after everything that could not survive a control was cut away. "
    "That is not the lesser way to reach a result; it is the way that produces one "
    "you can defend.")

# ============================================================ 5
h1("5.  The honest bottom line")
body(
    "I did not prove the grand thing the optimistic version kept promising — a "
    "powerful new kind of computer. I proved a smaller, exact thing: that point "
    "symmetry gates even-order computation in a physical resonator, with a clean "
    "rule for exactly which modes are silenced and which are not. The underlying "
    "mathematics (symmetry selection rules) is a known and respected tool; what is "
    "new is stating it cleanly in this setting, demonstrating it rigorously, "
    "correcting the framing from ‘aperiodicity’ to ‘symmetry,’ and "
    "mapping exactly where it does and does not bite. It is a modest, exact, "
    "well-bounded contribution — and the discipline of refusing to claim more "
    "than I could show, including against my own enthusiasm, is what makes it one "
    "worth standing behind.")

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Selection_Rule_Explained.docx")
doc.save(out)
print("Saved", out)
