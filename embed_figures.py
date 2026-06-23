"""Insert the actual figure images into PAPER_DRAFT.docx, right after the
paragraphs that reference each figure's filename."""
from docx import Document
from docx.shared import Inches
import copy

doc = Document("PAPER_DRAFT.docx")

targets = [
    ("figure_convergence.png", "figure_convergence.png", 5.0),
    ("figure_mode_shapes.png", "figure_mode_shapes.png", 6.0),
]

body = doc.element.body
paragraphs = list(doc.paragraphs)

for marker, img_path, width_in in targets:
    for p in paragraphs:
        if marker in p.text:
            # insert a new paragraph with the image right after this one
            new_p = doc.add_paragraph()
            run = new_p.add_run()
            run.add_picture(img_path, width=Inches(width_in))
            # move the new paragraph's XML to right after the marker paragraph
            p._p.addnext(new_p._p)
            break

doc.save("PAPER_DRAFT.docx")
print("Figures embedded.")
